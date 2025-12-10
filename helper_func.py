import re
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def  display_email_campaigns(gen_email_campaign_a, gen_email_campaign_b):
    if gen_email_campaign_a and gen_email_campaign_b:
        # Create a dictionary to format data into a table
        data = {
            "Attribute": ["Variant", "Subject", "Body", "Tone", "Call-to-Action (CTA)", "Target Audience"],
            "Campaign A": [
                gen_email_campaign_a.variant,
                gen_email_campaign_a.subject,
                gen_email_campaign_a.body,
                gen_email_campaign_a.tone,
                gen_email_campaign_a.call_to_action,
                gen_email_campaign_a.target_audience
            ],
            "Campaign B": [
                gen_email_campaign_b.variant,
                gen_email_campaign_b.subject,
                gen_email_campaign_b.body,
                gen_email_campaign_b.tone,
                gen_email_campaign_b.call_to_action,
                gen_email_campaign_b.target_audience
            ],
        }

        # Convert to Pandas DataFrame for better table formatting
        return pd.DataFrame(data)
    
def display_user_personas(users):
    dct_user_personas = {
        "persona_id":[],
        "name":[],
        "age":[],
        "gender":[],
        "occupation":[],
        "interests":[],
        "digital_behavior":[],
        "campaign_varient":[]
    }
    for user in users.personas:
        dct_user_personas["persona_id"].append(user.persona_id)
        dct_user_personas["name"].append(user.name)
        dct_user_personas["age"].append(user.age)
        dct_user_personas["gender"].append(user.gender)
        dct_user_personas["occupation"].append(user.occupation)
        dct_user_personas["interests"].append(user.interests)
        dct_user_personas["digital_behavior"].append(user.digital_behavior)
        dct_user_personas["campaign_varient"].append(user.campaign_varient)
    
    return pd.DataFrame(dct_user_personas)

def display_user_responses(responses):
    dct_user_response = {
        "person_id":[],
        "campaign_varient":[],
        "email_motiv_opened":[],
        "time_to_open_email":[],
        "ad_clicked":[],
        "time_to_click_ad":[],
        "email_review":[],
        "conversion":[]
    }

    for resp in responses:
        dct_user_response["person_id"].append(resp["id"])
        dct_user_response["campaign_varient"].append(resp["varient"])
        dct_user_response["email_motiv_opened"].append(resp["response"].email_motiv_opened)
        dct_user_response["time_to_open_email"].append(resp["response"].time_to_open_email)
        dct_user_response["ad_clicked"].append(resp["response"].ad_clicked)
        dct_user_response["time_to_click_ad"].append(resp["response"].time_to_click_ad)
        dct_user_response["email_review"].append(resp["response"].email_review)
        dct_user_response["conversion"].append(resp["response"].conversion)
    
    return pd.DataFrame(dct_user_response)

def save_as_docx(dct_text, filename, isWF=False):
    doc = Document()
    
    # -----------------------------------------------------------------
    # Helper function to process text (handles markdown, lists, and bold)
    # -----------------------------------------------------------------
    def add_formatted_content(doc_object, content, style=None):
        # 1. Split content by double newline to treat as distinct blocks/paragraphs
        blocks = content.split('\n\n')
        
        for block in blocks:
            block = block.strip()
            if not block:
                continue
                
            # 2. Check for bullet points (-, *, or •)
            if re.match(r'^\s*[-*•]\s', block, re.MULTILINE):
                # Process lists line by line
                list_lines = block.split('\n')
                for line in list_lines:
                    line = line.strip()
                    if re.match(r'^\s*[-*•]\s', line):
                        # Add as a bulleted list item, stripping the marker
                        doc_object.add_paragraph(
                            re.sub(r'^\s*[-*•]\s', '', line), 
                            style='List Bullet'
                        )
                    else:
                        # Add subsequent non-list lines in the same block as regular text
                        add_formatted_content(doc_object, line)
            
            # 3. Check for simple numbered lists (e.g., '1.', '2.)')
            elif re.match(r'^\s*\d+[\.\)]\s', block, re.MULTILINE):
                list_lines = block.split('\n')
                for line in list_lines:
                    line = line.strip()
                    if re.match(r'^\s*\d+[\.\)]\s', line):
                        # Add as a numbered list item, stripping the marker
                        doc_object.add_paragraph(
                            re.sub(r'^\s*\d+[\.\)]\s', '', line), 
                            style='List Number'
                        )
                    else:
                        add_formatted_content(doc_object, line)
                        
            # 4. If not a list, treat as a regular paragraph and check for bolding
            else:
                p = doc_object.add_paragraph('', style=style)
                
                # Split by ** to find bolded sections
                parts = re.split(r'(\*\*.*?\*\*)', block)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Add run with bold style
                        p.add_run(part.strip('**')).bold = True
                    else:
                        # Add run with normal style
                        p.add_run(part)

    # -----------------------------------------------------------------
    # REPORT STRUCTURE
    # -----------------------------------------------------------------
    
    # Title Page/Header
    doc.add_heading("A/B Test Report", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(" Email Campaign Analysis", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Map sections and titles
    sections = [
        ("1. Introduction", dct_text.Introduction),
        ("2. Experiment Process (Deep Dive)", dct_text.Experiment_process),
        ("3. Email Campaign Analysis (A vs. B)", dct_text.Email_Campaign_Analysis),
        ("4. User Persona Analysis & Segmentation", dct_text.User_Persona_Analysis),
        ("5. User Response Analysis (Qualitative)", dct_text.User_Response_Analysis),
        ("6. Performance Metrics Breakdown", dct_text.Performance_Metrics),
        ("7. Interpretations & Business Implications", dct_text.Interpretations),
        ("8. Recommendations for Next Steps", dct_text.Recommendations),
        ("9. Conclusion", dct_text.Conclusion)
    ]

    for title, content in sections:
        doc.add_heading(title, level=2)
        add_formatted_content(doc, content)
        doc.add_page_break() # Ensures each section starts on a new page

    if isWF:
        doc.save(filename)

    return doc

def display_experiment(response):
    # Create a clean container for the experiment
    with st.container(height=350):
        # 1. Header with ID
        st.subheader(f"🧪 Experiment ID: {response.experiment_id}")

        # 3. Experiment Guidelines with clear hierarchy
        st.markdown("### 📋 Experiment Guidelines")
    
        # Use a colored box or just clean markdown for the guidelines
        st.markdown(response.experiment_guidelines)
    
    return